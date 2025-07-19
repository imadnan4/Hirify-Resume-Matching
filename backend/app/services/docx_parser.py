"""
DOCX document parser using python-docx library with advanced features.
"""
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Optional, Any
import logging
import re
import os
from pathlib import Path
import zipfile
import xml.etree.ElementTree as ET

logger = logging.getLogger(__name__)


class DOCXParsingError(Exception):
    """Custom exception for DOCX parsing errors."""
    pass


class DOCXParser:
    """Service for parsing DOCX documents and extracting text content."""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.doc']
        
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX file with metadata and structure preservation.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dict containing extracted text and metadata
            
        Raises:
            DOCXParsingError: If DOCX parsing fails
        """
        try:
            result = {
                'text': '',
                'paragraphs': [],
                'tables': [],
                'metadata': {},
                'structure': {},
                'extraction_method': 'python-docx',
                'errors': [],
                'warnings': []
            }
            
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.docx':
                result = self._extract_from_docx(file_path)
            elif file_extension == '.doc':
                result = self._extract_from_doc(file_path)
            else:
                raise DOCXParsingError(f"Unsupported file format: {file_extension}")
            
            # Post-process the extracted text
            result = self._post_process_text(result)
            
            # Validate extraction quality
            self._validate_extraction_quality(result)
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from document {file_path}: {str(e)}")
            raise DOCXParsingError(f"Document parsing failed: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        result = {
            'text': '',
            'paragraphs': [],
            'tables': [],
            'metadata': {},
            'structure': {},
            'extraction_method': 'python-docx',
            'errors': [],
            'warnings': []
        }
        
        try:
            # Open the document
            doc = Document(file_path)
            
            # Extract metadata
            result['metadata'] = self._extract_metadata(doc)
            
            # Extract paragraphs
            all_text = []
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    all_text.append(text)
                    
                    # Extract paragraph-specific information
                    para_info = {
                        'paragraph_number': i + 1,
                        'text': text,
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'alignment': str(paragraph.alignment) if paragraph.alignment else 'None',
                        'char_count': len(text),
                        'word_count': len(text.split()),
                        'runs': []
                    }
                    
                    # Extract run information (formatting)
                    for run in paragraph.runs:
                        if run.text.strip():
                            run_info = {
                                'text': run.text,
                                'bold': run.bold,
                                'italic': run.italic,
                                'underline': run.underline,
                                'font_name': run.font.name if run.font.name else 'Default',
                                'font_size': run.font.size.pt if run.font.size else None
                            }
                            para_info['runs'].append(run_info)
                    
                    result['paragraphs'].append(para_info)
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_info = self._extract_table_data(table, i + 1)
                result['tables'].append(table_info)
                
                # Add table text to main text
                table_text = self._table_to_text(table_info)
                if table_text:
                    all_text.append(table_text)
            
            # Combine all text
            result['text'] = '\n\n'.join(all_text)
            
            # Extract document structure
            result['structure'] = self._analyze_document_structure(result['text'], result['paragraphs'])
            
            # Extract additional document properties
            result['metadata'].update(self._extract_document_properties(file_path))
            
        except Exception as e:
            result['errors'].append(f"Error extracting DOCX content: {str(e)}")
            logger.error(f"Error extracting DOCX content from {file_path}: {str(e)}")
        
        return result
    
    def _extract_from_doc(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOC file (legacy format).
        Note: This is limited compared to DOCX extraction.
        """
        result = {
            'text': '',
            'paragraphs': [],
            'tables': [],
            'metadata': {},
            'structure': {},
            'extraction_method': 'limited-doc',
            'errors': [],
            'warnings': ['DOC format has limited extraction capabilities']
        }
        
        try:
            # Try to use python-docx (may work for some DOC files)
            doc = Document(file_path)
            
            # Extract basic text
            all_text = []
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    all_text.append(text)
                    
                    para_info = {
                        'paragraph_number': i + 1,
                        'text': text,
                        'style': 'Unknown',
                        'alignment': 'None',
                        'char_count': len(text),
                        'word_count': len(text.split()),
                        'runs': []
                    }
                    
                    result['paragraphs'].append(para_info)
            
            result['text'] = '\n\n'.join(all_text)
            result['structure'] = self._analyze_document_structure(result['text'], result['paragraphs'])
            
        except Exception as e:
            result['errors'].append(f"Error extracting DOC content: {str(e)}")
            logger.error(f"Error extracting DOC content from {file_path}: {str(e)}")
            
            # If python-docx fails, we could try alternative methods here
            # For now, we'll just return the error
            
        return result
    
    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extract metadata from DOCX document."""
        metadata = {
            'author': '',
            'title': '',
            'subject': '',
            'keywords': '',
            'created': '',
            'modified': '',
            'last_modified_by': '',
            'category': '',
            'comments': '',
            'paragraphs_count': len(doc.paragraphs),
            'tables_count': len(doc.tables)
        }
        
        try:
            # Extract core properties
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    'author': core_props.author or '',
                    'title': core_props.title or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else '',
                    'last_modified_by': core_props.last_modified_by or '',
                    'category': core_props.category or '',
                    'comments': core_props.comments or ''
                })
        
        except Exception as e:
            logger.warning(f"Error extracting document metadata: {str(e)}")
        
        return metadata
    
    def _extract_document_properties(self, file_path: str) -> Dict[str, Any]:
        """Extract additional document properties from DOCX file."""
        properties = {
            'file_size': 0,
            'word_count': 0,
            'character_count': 0,
            'page_count': 0,
            'language': '',
            'application': ''
        }
        
        try:
            # Get file size
            properties['file_size'] = os.path.getsize(file_path)
            
            # Extract properties from the DOCX structure
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Try to read app.xml for additional properties
                try:
                    app_xml = docx_zip.read('docProps/app.xml')
                    app_root = ET.fromstring(app_xml)
                    
                    # Extract various properties
                    for elem in app_root:
                        tag = elem.tag.split('}')[-1]  # Remove namespace
                        if tag == 'Pages':
                            properties['page_count'] = int(elem.text) if elem.text else 0
                        elif tag == 'Words':
                            properties['word_count'] = int(elem.text) if elem.text else 0
                        elif tag == 'Characters':
                            properties['character_count'] = int(elem.text) if elem.text else 0
                        elif tag == 'Application':
                            properties['application'] = elem.text or ''
                
                except Exception as e:
                    logger.warning(f"Could not extract app.xml properties: {str(e)}")
                
                # Try to read core.xml for additional properties
                try:
                    core_xml = docx_zip.read('docProps/core.xml')
                    core_root = ET.fromstring(core_xml)
                    
                    for elem in core_root:
                        tag = elem.tag.split('}')[-1]  # Remove namespace
                        if tag == 'language':
                            properties['language'] = elem.text or ''
                
                except Exception as e:
                    logger.warning(f"Could not extract core.xml properties: {str(e)}")
        
        except Exception as e:
            logger.warning(f"Error extracting document properties: {str(e)}")
        
        return properties
    
    def _extract_table_data(self, table, table_number: int) -> Dict[str, Any]:
        """Extract data from a table."""
        table_info = {
            'table_number': table_number,
            'rows': len(table.rows),
            'columns': len(table.columns) if table.rows else 0,
            'data': [],
            'headers': []
        }
        
        try:
            # Extract table data
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                
                table_info['data'].append(row_data)
                
                # Assume first row is headers
                if row_idx == 0:
                    table_info['headers'] = row_data
        
        except Exception as e:
            logger.warning(f"Error extracting table data: {str(e)}")
        
        return table_info
    
    def _table_to_text(self, table_info: Dict[str, Any]) -> str:
        """Convert table data to readable text."""
        if not table_info['data']:
            return ''
        
        text_lines = []
        text_lines.append(f"[Table {table_info['table_number']}]")
        
        for row in table_info['data']:
            # Join cells with tabs
            text_lines.append('\t'.join(row))
        
        return '\n'.join(text_lines)
    
    def _analyze_document_structure(self, text: str, paragraphs: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Analyze document structure and identify sections."""
        structure = {
            'sections': [],
            'has_contact_info': False,
            'has_experience': False,
            'has_education': False,
            'has_skills': False,
            'estimated_sections': [],
            'heading_paragraphs': [],
            'formatted_text': {
                'bold_paragraphs': [],
                'italic_paragraphs': [],
                'underlined_paragraphs': []
            }
        }
        
        if not text:
            return structure
        
        # Common section patterns in resumes
        section_patterns = {
            'contact': r'(?i)(contact|phone|email|address|linkedin)',
            'summary': r'(?i)(summary|profile|objective|about)',
            'experience': r'(?i)(experience|employment|work\s+history|professional)',
            'education': r'(?i)(education|academic|degree|university|college)',
            'skills': r'(?i)(skills|technical|competencies|abilities)',
            'projects': r'(?i)(projects|portfolio)',
            'certifications': r'(?i)(certifications|certificates|licenses)',
            'achievements': r'(?i)(achievements|awards|accomplishments)'
        }
        
        # Analyze text for different sections
        for section_name, pattern in section_patterns.items():
            if re.search(pattern, text):
                structure[f'has_{section_name}'] = True
                structure['estimated_sections'].append(section_name)
        
        # Analyze paragraph formatting to identify structure
        for para in paragraphs:
            # Check for heading styles
            if 'heading' in para.get('style', '').lower():
                structure['heading_paragraphs'].append(para)
            
            # Check for formatted text
            for run in para.get('runs', []):
                if run.get('bold'):
                    structure['formatted_text']['bold_paragraphs'].append(para['paragraph_number'])
                if run.get('italic'):
                    structure['formatted_text']['italic_paragraphs'].append(para['paragraph_number'])
                if run.get('underline'):
                    structure['formatted_text']['underlined_paragraphs'].append(para['paragraph_number'])
        
        # Try to identify section boundaries based on formatting and content
        current_section = None
        
        for para in paragraphs:
            para_text = para['text']
            
            # Check if paragraph looks like a section header
            if self._is_section_header(para_text, para):
                if current_section:
                    structure['sections'].append(current_section)
                
                current_section = {
                    'title': para_text,
                    'content': '',
                    'paragraph_count': 0,
                    'formatting': {
                        'style': para.get('style', ''),
                        'has_bold': any(run.get('bold') for run in para.get('runs', [])),
                        'has_italic': any(run.get('italic') for run in para.get('runs', [])),
                        'has_underline': any(run.get('underline') for run in para.get('runs', []))
                    }
                }
            elif current_section:
                current_section['content'] += para_text + '\n'
                current_section['paragraph_count'] += 1
        
        # Add the last section
        if current_section:
            structure['sections'].append(current_section)
        
        return structure
    
    def _is_section_header(self, text: str, para: Dict[str, Any]) -> bool:
        """Determine if a paragraph looks like a section header."""
        # Check length - headers are usually short
        if len(text) > 50:
            return False
        
        # Check for heading styles
        if 'heading' in para.get('style', '').lower():
            return True
        
        # Check for formatting (bold, all caps, etc.)
        has_bold = any(run.get('bold') for run in para.get('runs', []))
        has_underline = any(run.get('underline') for run in para.get('runs', []))
        
        if has_bold or has_underline:
            return True
        
        # Check for common header patterns
        header_patterns = [
            r'^[A-Z\s]+$',  # All caps
            r'^[A-Za-z\s]+:$',  # Ends with colon
            r'^\d+\.\s+[A-Za-z]',  # Numbered section
            r'^[A-Z][a-z]+\s+[A-Z][a-z]+$'  # Title case
        ]
        
        for pattern in header_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def _post_process_text(self, result: Dict[str, Any]) -> Dict[str, Any]:
        """Post-process extracted text to improve quality."""
        if not result['text']:
            return result
        
        # Clean up text
        text = result['text']
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Replace multiple newlines with double newline
        text = re.sub(r' +', ' ', text)  # Replace multiple spaces with single space
        
        # Fix common extraction issues
        text = text.replace('\u2022', '•')  # Replace bullet points
        text = text.replace('\u2013', '-')  # Replace en dash
        text = text.replace('\u2014', '--')  # Replace em dash
        
        # Clean up tabs from tables
        text = re.sub(r'\t+', '\t', text)  # Replace multiple tabs with single tab
        
        result['text'] = text
        
        # Update statistics
        result['statistics'] = {
            'total_characters': len(result['text']),
            'total_words': len(result['text'].split()),
            'total_lines': len(result['text'].split('\n')),
            'total_paragraphs': len(result['paragraphs']),
            'total_tables': len(result['tables'])
        }
        
        return result
    
    def _validate_extraction_quality(self, result: Dict[str, Any]) -> None:
        """Validate the quality of text extraction."""
        if not result['text']:
            result['warnings'].append("No text was extracted from the document")
            return
        
        text = result['text']
        
        # Check for minimum content
        if len(text.split()) < 10:
            result['warnings'].append("Very little text was extracted")
        
        # Check for reasonable structure
        if len(result['paragraphs']) < 2:
            result['warnings'].append("Document appears to have very few paragraphs")
        
        # Check for garbled text (high ratio of non-alphabetic characters)
        alpha_chars = sum(1 for char in text if char.isalpha())
        total_chars = len(text.replace(' ', '').replace('\n', '').replace('\t', ''))
        
        if total_chars > 0 and alpha_chars / total_chars < 0.7:
            result['warnings'].append("Extracted text may contain garbled characters")
        
        logger.info(f"Text extraction quality check completed. Characters: {len(text)}, Words: {len(text.split())}")
    
    def is_docx_file(self, file_path: str) -> bool:
        """Check if file is a valid DOCX file."""
        try:
            doc = Document(file_path)
            return True
        except:
            return False
    
    def get_paragraph_count(self, file_path: str) -> int:
        """Get the number of paragraphs in a DOCX file."""
        try:
            doc = Document(file_path)
            return len(doc.paragraphs)
        except:
            return 0
