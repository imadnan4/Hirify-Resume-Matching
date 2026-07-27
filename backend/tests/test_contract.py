from __future__ import annotations

from io import BytesIO

import fitz
from docx import Document


def build_resume_pdf_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page()
    text = """
Jane Doe
Karachi, Pakistan
jane@example.com
+92 300 1234567
https://www.linkedin.com/in/janedoe

Summary
Frontend engineer with React and TypeScript experience building production dashboards.

Work Experience
Senior Frontend Engineer at Acme
Jan 2021 - Present
Built React, TypeScript, GraphQL, and Tailwind interfaces for hiring workflows.

Education
Bachelor of Science in Computer Science
NED University
2020

Skills
React, TypeScript, GraphQL, Tailwind, CSS, HTML, Git

Certifications
AWS Certified Cloud Practitioner
""".strip()
    page.insert_text((72, 72), text, fontsize=11)
    data = document.tobytes()
    document.close()
    return data


def build_resume_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("John Smith")
    doc.add_paragraph("john@example.com")
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Backend engineer with Python, FastAPI, and PostgreSQL experience.")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, FastAPI, PostgreSQL, Docker")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def upload_resume(client, filename: str, content: bytes, content_type: str):
    return client.post(
        "/api/v1/resumes/upload",
        files={"file": (filename, content, content_type)},
    )


def test_health_endpoint(client):
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "healthy", "message": "API is running"}


def test_resume_upload_preview_and_reprocess_contract(client):
    response = upload_resume(client, "resume.pdf", build_resume_pdf_bytes(), "application/pdf")
    assert response.status_code == 201, response.text
    payload = response.json()
    resume_id = payload["id"]
    assert payload["filename"] == "resume.pdf"
    assert payload["status"] == "completed"

    list_response = client.get("/api/v1/resumes/", params={"limit": 1000})
    assert list_response.status_code == 200
    listed = list_response.json()
    assert listed["total"] == 1
    assert listed["items"][0]["status"] == "completed"

    preview_response = client.get(f"/api/v1/resumes/{resume_id}/preview")
    assert preview_response.status_code == 200
    preview = preview_response.json()
    assert isinstance(preview["skills"], list)
    assert "react" in [item.lower() for item in preview["skills"]]
    assert isinstance(preview["work_experience"], list)
    assert isinstance(preview["education"], list)
    assert isinstance(preview["certifications"], list)
    assert preview["contact_info"]["email"] == "jane@example.com"

    status_response = client.get(f"/api/v1/resumes/{resume_id}/status")
    assert status_response.status_code == 200
    assert status_response.json()["progress"] == 100

    reprocess_response = client.post(f"/api/v1/resumes/{resume_id}/reprocess")
    assert reprocess_response.status_code == 200
    assert reprocess_response.json()["message"] == "Resume reprocessing started"


def test_docx_supported_and_doc_rejected(client):
    docx_response = upload_resume(
        client,
        "resume.docx",
        build_resume_docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert docx_response.status_code == 201, docx_response.text
    preview_response = client.get(f"/api/v1/resumes/{docx_response.json()['id']}/preview")
    assert preview_response.status_code == 200, preview_response.text
    preview = preview_response.json()
    assert isinstance(preview["skills"], list)
    assert "python" in [item.lower() for item in preview["skills"]]

    doc_response = client.post(
        "/api/v1/resumes/upload",
        files={"file": ("legacy.doc", b"fake-binary", "application/msword")},
    )
    assert doc_response.status_code == 400
    assert "Legacy .doc files are not supported" in doc_response.json()["detail"]


def test_jobs_and_matching_contract_many_to_many_bulk(client):
    resume1_response = upload_resume(client, "resume-a.pdf", build_resume_pdf_bytes(), "application/pdf")
    assert resume1_response.status_code == 201, resume1_response.text
    resume1 = resume1_response.json()
    resume2_response = upload_resume(
        client,
        "resume-b.docx",
        build_resume_docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert resume2_response.status_code == 201, resume2_response.text
    resume2 = resume2_response.json()

    job1_response = client.post(
        "/api/v1/jobs/",
        json={
            "title": "Frontend Engineer",
            "company": "Acme",
            "description": "Build React and TypeScript interfaces for hiring software.",
            "requirements": "3+ years experience with React, TypeScript, GraphQL and CSS.",
            "source": "manual",
            "location": "Remote",
            "employment_type": "full-time",
            "experience_level": "mid",
        },
    )
    assert job1_response.status_code == 201, job1_response.text
    job1 = job1_response.json()
    assert isinstance(job1["extracted_skills"], list)

    job2_response = client.post(
        "/api/v1/jobs/",
        json={
            "title": "Backend Engineer",
            "company": "Beta",
            "description": "Build APIs using Python and PostgreSQL.",
            "requirements": "2+ years with Python, FastAPI, Docker and SQL.",
            "source": "manual",
        },
    )
    assert job2_response.status_code == 201, job2_response.text
    job2 = job2_response.json()

    single_match = client.post(
        "/api/v1/matching/match",
        json={"resume_id": resume1["id"], "job_id": job1["id"]},
    )
    assert single_match.status_code == 201, single_match.text
    single_payload = single_match.json()
    assert 0 <= single_payload["overall_score"] <= 1

    repeat_single = client.post(
        "/api/v1/matching/match",
        json={"resume_id": resume1["id"], "job_id": job1["id"]},
    )
    assert repeat_single.status_code == 201

    bulk_match = client.post(
        "/api/v1/matching/bulk-match",
        json={
            "resume_ids": [resume1["id"], resume2["id"]],
            "job_ids": [job1["id"], job2["id"]],
            "min_score_threshold": 0.0,
            "include_explanations": True,
        },
    )
    assert bulk_match.status_code == 200, bulk_match.text
    bulk_payload = bulk_match.json()
    assert bulk_payload["total_matches"] == 4
    assert len(bulk_payload["matches"]) == 4

    matches_response = client.get("/api/v1/matching/", params={"limit": 1000})
    assert matches_response.status_code == 200
    matches = matches_response.json()
    assert matches["total"] == 4
    assert all(0 <= item["overall_score"] <= 1 for item in matches["items"])

    stats_response = client.get("/api/v1/matching/stats")
    assert stats_response.status_code == 200
    assert stats_response.json()["total_matches"] == 4

    candidates_response = client.get(f"/api/v1/matching/job/{job1['id']}/candidates")
    assert candidates_response.status_code == 200
    assert len(candidates_response.json()) >= 1

    delete_match_id = matches["items"][0]["id"]
    delete_match_response = client.delete(f"/api/v1/matching/{delete_match_id}")
    assert delete_match_response.status_code == 200

    delete_resume_response = client.delete(f"/api/v1/resumes/{resume2['id']}")
    assert delete_resume_response.status_code == 200

    delete_job_response = client.delete(f"/api/v1/jobs/{job2['id']}")
    assert delete_job_response.status_code == 200
